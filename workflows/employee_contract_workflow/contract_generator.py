from abstra.forms import *
from abstra.tables import run
import abstra.workflows as aw
from datetime import datetime, date
import os
from unidecode import unidecode
from docxtpl import DocxTemplate
import re
from docx import Document
import pathlib

# Checking the email domain
user = get_user()
if not user.email.endswith("@abstra.app"):
    display("Unauthorized access. Please contact admin@abstra.app.")
    exit()

# In this code we are going to set a default document format with some tags to be filled
# the tags are defined by {{tag_name}}
# the document is going to be automatically filled with the user's data
# the user can also upload a model of a document and fill it with the tags

DATA_WORKSHEET = "Data"
FILES_BASE_PATH = "/tmp/fill_my_docs"
REGEX = r"\{\{(.*?)\}\}"


def get_team_info(team_id):
    sql = "SELECT salary, complement_address, position, bank_account_number, id_emited_by, taxpayer_id, address, birth_date, zip_code, bank_name, email, number_address, phone_number, abstra_email, name, bank_branch_code, started_at, district, country, identification_number FROM team WHERE id = $1;"
    params = [team_id]
    return run(sql, params)[0]


# the tag cant contain " " so whe have to replace it with "_"
# and after the user fills the tag, we have to replace it again with " "
def transform_tags(tags):
    transformed = []
    for tag in tags:
        transformed.append(
            {
                "value": unidecode((tag.lower())),
                "label": tag.replace("_", " ").capitalize(),
            }
        )
    return transformed


def create_new_doc_with_tags(tags_values_dict, filepath, filename):
    doc = DocxTemplate(filepath)
    context = tags_values_dict
    try:
        doc.render(context)
    except Exception as e:
        problematic_tags = [k for k in context.keys() if k in str(e)]
        display(f"Error: {e}. Please check the following tags: {problematic_tags}")
    doc.save(filepath)
    output_filepath = os.getcwd() + "/" + filepath
    return output_filepath


def render(partial):
    if len(partial) != 0:
        if partial.get("existing_contract") == "New Contract":
            return Page().read_file("Upload your contract", key="contract")


def generate_document(
    file_response, team_member_name, contract_folder, contract_data={}
):
    filename = f"{date.today().strftime('%Y%m%d')}_{team_member_name}"
    filename += ".docx"
    filepath = os.path.join(contract_folder, f"{filename}.docx")
    with open(filepath, "wb") as out_file:
        out_file.write(file_response)
    doc = Document(filepath)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    tags = re.findall(REGEX, text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_tags = re.findall(REGEX, cell.text)
                if cell_tags:
                    tags.append(cell_tags[0])
    tags_dict = dict.fromkeys(tags)
    tags_original_list = list(tags_dict.keys())
    tags_original = [t.strip() for t in tags_original_list]
    new_tags_original = [
        tag
        for tag in tags_original
        if tag not in list(contract_data.keys()) or contract_data[tag] is None
    ]
    tags_formatted = transform_tags(new_tags_original)
    # for inputs in tags_with_underscores:
    page = Page()
    for tags in tags_formatted:
        tag = tags["label"]
        page = page.read(f"{tag}:", placeholder=tag)
    response = page.run()
    response_list = list(response.values())
    tags_response = dict(zip(new_tags_original, response_list))
    tags_response.update(contract_data)
    output_filepath = create_new_doc_with_tags(tags_response, filepath, filename)
    display_file(
        output_filepath, download_text="Click here to download the filled document"
    )
    return output_filepath


def render(partial):
    if len(partial) != 0:
        if partial["existing_contract"] == "New Contract":
            return Page().read_file("Upload your contract", key="contract")


contract = (
    Page()
    .read_dropdown(
        "Contract Model",
        ["New Contract", "Individual"],
        key="existing_contract",
    )
    # Will render depending of what answer you give in the dropdown
    .reactive(render)
    .run()
)
folder = pathlib.Path(os.environ.get("ABSTRA_FILES_FOLDER", "./"))
contract_folder = folder / "contracts"
contract_folder.mkdir(parents=True, exist_ok=True)
stage = aw.get_stage()
team_id = stage["id"]
name = stage["name"]
email = stage["email"]
id_taxpayer = stage["id_taxpayer"]
team_info = get_team_info(team_id)
(
    salary,
    complement_address,
    position,
    bank_account_number,
    id_emited_by,
    taxpayer_id,
    address,
    birth_date,
    zip_code,
    bank_name,
    email,
    number_address,
    phone_number,
    abstra_email,
    name,
    bank_branch_code,
    started_at,
    district,
    country,
    identification_number,
) = team_info.values()
if complement_address is None:
    complement_address = ""
if contract["existing_contract"] == "New Contract":
    if not contract["contract"].file.name.endswith(".docx"):
        display("Please upload a .docx file")
        exit()
    contract = contract["contract"].file.read()
    output_filepath = generate_document(contract, name, contract_folder)
else:
    months = [
        "Janeiro",
        "Fevereiro",
        "Março",
        "Abril",
        "Maio",
        "Junho",
        "Julho",
        "Agosto",
        "Setembro",
        "Outubro",
        "Novembro",
        "Dezembro",
    ]
    start_abstra_at = datetime.strptime(started_at, "%Y-%m-%dT%H:%M:%S.000Z")
    document_filename = "Contrato de Trabalho"
    contract = open(f"contract_models/individual_contract_ptbr.docx", "rb").read()
    contract_data = {
        "nome_completo_contratado": name,
        "nacionalidade_contratado": country,
        "estado_civil_contratado": "preencher",
        "profissao_contratado": position,
        "cpf_contratado": identification_number,
        "rg_contratado": "preencher",
        "emissao_por_rg": id_emited_by,
        "endereco_completo_contratado": address,
        "complemento_endereco_contratado": complement_address,
        "bairro_contratado": district,
        "cep_contratado": zip_code,
        "email_contratado": email,
        "carga_horaria_contratado": "preencher",
        "carga_horaria_contratado_extenso": "preencher",
        "carga_horaria_semamal_contratado": "preencher",
        "carga_semanal_horaria_contratado": "preencher",
        "remuneracao_contratado": salary,
        "remuneracao_extenso_contratado": "preencher",
        "banco_contratado": bank_name,
        "agencia_contratado": bank_branch_code,
        "conta_corrente_contratado": bank_account_number,
        "pix_contratado": "preencher",
        "regime_de_trabalho_contratado": "preencher",
        "data_assinatura_contratado": f"{start_abstra_at.day} de {months[start_abstra_at.month - 1]} de {start_abstra_at.year}",
    }
    document_filename = "Contrato de Prestação de Serviços e Outras Avenças"
    output_filepath = generate_document(contract, name, contract_folder, contract_data)
aw.next_stage(
    [
        {
            "assignee": email,
            "data": {
                "id": team_id,
                "name": name,
                "email": email,
                "output_filepath": output_filepath,
                "document_filename": document_filename,
                "id_taxpayer": id_taxpayer,
            },
            "stage": "contract-approval",
        }
    ]
)
