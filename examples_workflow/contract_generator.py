from abstra.forms import *
import abstra.workflows as aw
from datetime import datetime, date
import os
from unidecode import unidecode
from docxtpl import DocxTemplate
import re
from docx import Document

import pathlib

user = get_user()
if not user.email.endswith("@abstra.app"):
    display("Unauthorized access. Please contact admin@abstra.app.")
    exit()


DATA_WORKSHEET = "Data"
FILES_BASE_PATH = "/tmp/fill_my_docs"
REGEX = r"\{\{(.*?)\}\}"

# Cria um dicionário com (variável, novo valor) para cada variável do contrato


def transform_tags(tags):
    transformed = []
    for tag in tags:
        transformed.append(
            {
                "value": unidecode((tag.lower())),
                "label": tag.replace("_", " ").capitalize(),
            }
        )
    return transformed


def create_new_doc_with_tags(tags_values_dict, filepath, filename):
    print("filepath", filepath)
    doc = DocxTemplate(filepath)
    context = tags_values_dict
    try:
        doc.render(context)
    except Exception as e:
        problematic_tags = [k for k in context.keys() if k in str(e)]
        display(f"Error: {e}. Please check the following tags: {problematic_tags}")

    doc.save(filepath)
    output_filepath = os.getcwd() + "/" + filepath
    return output_filepath


def generate_document(
    file_response, team_member_name, contract_folder, contract_data={}
):
    filename = f"{date.today().strftime('%Y%m%d')}_{team_member_name}"
    filename += ".docx"

    filepath = os.path.join(contract_folder, f"{filename}.docx")

    with open(filepath, "wb") as out_file:
        out_file.write(file_response)

    doc = Document(filepath)

    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    tags = re.findall(REGEX, text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_tags = re.findall(REGEX, cell.text)
                if cell_tags:
                    tags.append(cell_tags[0])

    tags_dict = dict.fromkeys(tags)
    tags_original_list = list(tags_dict.keys())
    tags_original = [t.strip() for t in tags_original_list]

    new_tags_original = [
        tag
        for tag in tags_original
        if tag not in list(contract_data.keys()) or contract_data[tag] is None
    ]

    tags_formatted = transform_tags(new_tags_original)

    # for inputs in tags_with_underscores:
    page = Page()
    for tags in tags_formatted:
        tag = tags["label"]
        page = page.read(f"{tag}:", placeholder=tag)

    response = page.run()
    response_list = list(response.values())

    tags_response = dict(zip(new_tags_original, response_list))

    tags_response.update(contract_data)

    output_filepath = create_new_doc_with_tags(tags_response, filepath, filename)

    display_file(
        output_filepath, download_text="Click here to download the filled document"
    )

    return output_filepath


contract_type = [
    {"label": "Individual", "value": "individual"},
    {"label": "Individual Freelancer", "value": "individual_freelancer"},
    {"label": "Company", "value": "company"},
]


def render(partial):
    if len(partial) != 0:
        if partial["existing_contract"] == "New Contract":
            return Page().read_file("Upload your contract", key="contract")
        else:
            return Page().read_dropdown(
                "Contract Type", contract_type, key="contract_type"
            )


contract = (
    Page()
    .read_dropdown(
        "Contract Model",
        ["New Contract", "Existing Contract model"],
        key="existing_contract",
    )
    .reactive(render)
    .run()
)

folder = pathlib.Path(os.environ.get("ABSTRA_FILES_FOLDER", "./"))
contract_folder = folder / "contracts"
contract_folder.mkdir(parents=True, exist_ok=True)


stage = aw.get_stage()
team_id = stage["id"]
name = stage["name"]
email = stage["email"]
id_taxpayer = stage["id_taxpayer"]
# team_id = 31
# name = "catarina@abstra.app"
# email = "Gabriel Saboya Nogueira Lima"

team_info = run_finance(
    'SELECT t.created_at, t.name, t.position, t.email, t.birth_date, t.identification_number, \
                    t.id_emited_by, t.address, t.number_address, t.complement_address, t.city, t.state,  t.district,\
                    t.zip_code, t.country, t.taxpayer_id, te.name as company_name, te.entity_number, te.district as company_district,\
                    te.state as company_state, te.number_address as company_number_address, te.complement_address as company_complement_address,\
                    te.address as company_address, te.zip_code as company_zip_code, te.city as company_city, tb.name as bank_name,\
                    tb.number as bank_number_account, tb.branch_code FROM "team" t LEFT JOIN "team_entity_informations" te\
                    ON t.id = te.team_id LEFT JOIN "team_bank_account" tb ON te.team_id = tb.team_id WHERE t.id = $1',
    params=[team_id],
)[0]

if team_info["complement_address"] is None:
    team_info["complement_address"] = ""

if team_info["company_complement_address"] is None:
    team_info["company_complement_address"] = ""

if contract["existing_contract"] == "New Contract":
    if not contract["contract"].file.name.endswith(".docx"):
        display("Please upload a .docx file")
        exit()
    contract = contract["contract"].file.read()
    output_filepath = generate_document(
        contract, team_info["company_name"], contract_folder
    )

else:
    months = [
        "Janeiro",
        "Fevereiro",
        "Março",
        "Abril",
        "Maio",
        "Junho",
        "Julho",
        "Agosto",
        "Setembro",
        "Outubro",
        "Novembro",
        "Dezembro",
    ]
    start_abstra_at = datetime.strptime(
        team_info["created_at"], "%Y-%m-%dT%H:%M:%S.000Z"
    )

    contract_type = contract["contract_type"]
    if contract_type == "individual":
        document_filename = "Contrato de Trabalho"

    elif contract_type == "individual_freelancer":
        contract = open(
            f"contract_models/{contract_type}_contract_ptbr.docx", "rb"
        ).read()

        contract_data = {
            "nome_completo_contratado": team_info["name"],
            "cpf_contratado": team_info["taxpayer_id"],
            "rg_contratado": team_info["identification_number"],
            "orgao_emissor_rg_contratado": team_info["id_emited_by"],
            "endereco_contratado": f'{team_info["address"]}, {team_info["number_address"]}',
            "complemento_contratado": f'{team_info["complement_address"]},',
            "bairro_contratado": team_info["district"],
            "cidade_contratado": team_info["city"],
            "cep_contratado": team_info["zip_code"],
            "email_contratado": team_info["email"],
            "nome_banco": team_info["bank_name"],
            "agencia_bancaria": team_info["branch_code"],
            "conta_corrente_bancaria": team_info["bank_number_account"],
            "data_inicio_na_abstra": f"{start_abstra_at.day} de {months[start_abstra_at.month - 1]} de {start_abstra_at.year}",
        }

        document_filename = "Contrato de Prestação de Serviços e Outras Avenças"

        output_filepath = generate_document(
            contract, team_info["name"], contract_folder, contract_data
        )

    else:
        contract = open(
            f"contract_models/{contract_type}_contract_ptbr.docx", "rb"
        ).read()

        contract_data = {
            "razao_social_da_contratada": team_info["company_name"],
            "cnpj_da_contratada": team_info["entity_number"],
            "endereco_da_contratada": f'{team_info["company_address"]}, {team_info["company_number_address"]}',
            "bairro_da_contratada": team_info["company_district"],
            "cep_da_contratada": team_info["company_zip_code"],
            "cidade_da_contratada": team_info["company_city"],
            "uf_da_contratada": team_info["company_state"],
            "email_do_contratado": team_info["email"],
            "nome_completo_do_contratado": team_info["name"],
            "rg_do_contratado": team_info["identification_number"],
            "orgao_emissor_rg_do_contratado": team_info["id_emited_by"],
            "cpf_do_contratado": team_info["taxpayer_id"],
            "data_inicio_abstra": f"{start_abstra_at.day} de {months[start_abstra_at.month - 1]} de {start_abstra_at.year}",
            "banco": team_info["bank_name"],
            "agencia": team_info["branch_code"],
            "conta_corrente": team_info["bank_number_account"],
        }

        document_filename = "Contrato de Prestação de Serviços"

        output_filepath = generate_document(
            contract, team_info["company_name"], contract_folder, contract_data
        )


aw.next_stage(
    [
        {
            "assignee": email,
            "data": {
                "id": team_id,
                "name": name,
                "email": email,
                "output_filepath": output_filepath,
                "document_filename": document_filename,
                "id_taxpayer": id_taxpayer,
            },
            "stage": "contract-approval",
        }
    ]
)
